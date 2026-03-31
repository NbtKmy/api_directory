from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── helper ──────────────────────────────────────────────────────────────────

def clean(text):
    return (text or "").replace("\xa0", " ").replace("\u202f", " ").strip()

def is_api_name(para):
    """API names are Normal paragraphs whose runs have a colored (non-None) font color."""
    if para.style.name not in ("Normal", "Default Paragraph Style"):
        return False
    colored = [r for r in para.runs if r.text.strip()
               and r.font.color and r.font.color.type is not None]
    return len(colored) > 0

def parse_doc(path):
    """
    Returns a list of dicts per API:
      name, url, description, table_rows (dict label→value)
    """
    doc = Document(path)
    body = doc.element.body
    apis = []
    current = None

    for child in body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para = next((p for p in doc.paragraphs if p._element is child), None)
            if para is None:
                continue
            text = clean(para.text)
            if not text:
                continue
            if para.style.name in ("Heading 1", "Heading 2", "Heading 3"):
                continue

            if text.startswith("http"):
                if current is not None:
                    current["url"] = text
            elif is_api_name(para):
                current = {"name": text, "url": "", "description": [], "table_rows": {}}
                apis.append(current)
            else:
                if current is not None:
                    current["description"].append(text)

        elif tag == "tbl":
            tbl = next((t for t in doc.tables if t._element is child), None)
            if tbl is None:
                continue
            if current is None:
                continue
            rows = {}
            for row in tbl.rows:
                cells = [clean(c.text) for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    rows[cells[0]] = cells[1]
            current["table_rows"] = rows

    for api in apis:
        api["description"] = " ".join(api["description"])

    return apis

# ── styling ──────────────────────────────────────────────────────────────────

HEADER_FILL = PatternFill("solid", fgColor="2E4057")
ALT_FILL    = PatternFill("solid", fgColor="E8EEF4")
NAME_FILL   = PatternFill("solid", fgColor="4A7BA6")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
NAME_FONT   = Font(bold=True, color="FFFFFF", size=11)
BODY_FONT   = Font(size=10)
thin        = Side(style="thin", color="CCCCCC")
BORDER      = Border(left=thin, right=thin, top=thin, bottom=thin)

def apply_header(ws, columns):
    for col_idx, name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=name)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
        cell.border = BORDER
    ws.row_dimensions[1].height = 30

def build_sheet(ws, apis, columns, field_keys, col_widths):
    apply_header(ws, columns)
    row = 2
    for i, api in enumerate(apis):
        # API name spans all columns
        n_cols = len(columns)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=n_cols)
        cell = ws.cell(row=row, column=1, value=api["name"])
        cell.fill = NAME_FILL
        cell.font = NAME_FONT
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        cell.border = BORDER
        ws.row_dimensions[row].height = 20
        row += 1

        # data row
        alt = (i % 2 == 1)
        fill = ALT_FILL if alt else PatternFill("solid", fgColor="FFFFFF")
        for col_idx, key in enumerate(field_keys, 1):
            if key == "url":
                value = api.get("url", "")
            elif key == "description":
                value = api.get("description", "")
            else:
                value = api["table_rows"].get(key, "")
            cell = ws.cell(row=row, column=col_idx, value=value)
            cell.fill = fill
            cell.font = BODY_FONT
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = BORDER
        ws.row_dimensions[row].height = 15
        row += 1

    for col_idx, width in col_widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    ws.freeze_panes = "A2"

# ── main ────────────────────────────────────────────────────────────────────

de_apis = parse_doc("API-Übersicht-DE.docx")
en_apis = parse_doc("API-Übersicht-EN.docx")

wb = openpyxl.Workbook()

# Sheet 1: Deutsch
ws_de = wb.active
ws_de.title = "Deutsch"
de_cols = ["URL", "Beschreibung", "Lizenz", "Inhalte & Formate", "Einschränkungen", "Dokumentation", "Antrag auf Zugang"]
de_keys = ["url", "description", "Lizenz", "Inhalte & Formate", "Einschränkungen", "Dokumentation", "Antrag auf Zugang"]
de_widths = {1: 40, 2: 35, 3: 35, 4: 55, 5: 40, 6: 40, 7: 40}
build_sheet(ws_de, de_apis, de_cols, de_keys, de_widths)

# Sheet 2: English
ws_en = wb.create_sheet(title="English")
en_cols = ["URL", "Description", "License", "Content & Formats", "Limitations", "Documentation", "How to Request Access"]
en_keys = ["url", "description", "License", "Content & formats", "Limitations", "Documentation", "How to request access"]
en_widths = {1: 40, 2: 35, 3: 35, 4: 55, 5: 40, 6: 40, 7: 40}
build_sheet(ws_en, en_apis, en_cols, en_keys, en_widths)

out = "API-Übersicht.xlsx"
wb.save(out)
print(f"Saved: {out}")
print(f"DE sheet: {len(de_apis)} APIs")
print(f"EN sheet: {len(en_apis)} APIs")
